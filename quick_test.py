from docx import Document

# Create a simple test document
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph.')
doc.save('test.docx')

# Try conversion
from docx2pdf import convert
convert('test.docx')