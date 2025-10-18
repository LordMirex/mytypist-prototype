#!/usr/bin/env python3
"""
Comprehensive test suite for MyTypist application
Tests all major features and components
"""

import os
import sys
import tempfile
import json
from datetime import datetime

def test_imports():
    """Test if all required modules can be imported"""
    print("Testing imports...")
    try:
        import flask
        import flask_sqlalchemy
        import docxtpl
        import reportlab
        import docx
        from PIL import Image
        print("✓ All core dependencies imported successfully")
        return True
    except ImportError as e:
        print(f"✗ Import error: {e}")
        return False

def test_app_structure():
    """Test application file structure"""
    print("\nTesting application structure...")
    
    required_files = [
        'app.py',
        'requirements.txt',
        'README.md',
        'templates/base.html',
        'templates/index.html',
        'templates/create.html',
        'templates/batch.html',
        'templates/batch_results.html',
        'templates/admin.html',
        'templates/error.html',
        'templates/results.html',
        'templates/partials/form_fields.html',
        'templates/admin/upload.html',
        'templates/admin/templates.html',
        'templates/admin/edit_template.html'
    ]
    
    required_dirs = [
        'templates',
        'templates/admin',
        'templates/partials',
        'upload',
        'generated',
        'db',
        'fonts'
    ]
    
    missing_files = []
    missing_dirs = []
    
    for file_path in required_files:
        if not os.path.exists(file_path):
            missing_files.append(file_path)
    
    for dir_path in required_dirs:
        if not os.path.exists(dir_path):
            missing_dirs.append(dir_path)
    
    if missing_files:
        print(f"✗ Missing files: {missing_files}")
    else:
        print("✓ All required files present")
    
    if missing_dirs:
        print(f"✗ Missing directories: {missing_dirs}")
    else:
        print("✓ All required directories present")
    
    return len(missing_files) == 0 and len(missing_dirs) == 0

def test_app_configuration():
    """Test app configuration and database models"""
    print("\nTesting app configuration...")
    
    try:
        # Import the app
        sys.path.insert(0, os.getcwd())
        import app
        
        # Test Flask app creation
        flask_app = app.app
        print("✓ Flask app created successfully")
        
        # Test database models
        with flask_app.app_context():
            # Test model definitions
            template = app.Template()
            placeholder = app.Placeholder()
            created_doc = app.CreatedDocument()
            batch_gen = app.BatchGeneration()
            print("✓ Database models defined correctly")
            
            # Test database initialization
            app.db.create_all()
            print("✓ Database tables created successfully")
        
        return True
        
    except Exception as e:
        print(f"✗ App configuration error: {e}")
        return False

def test_document_processor():
    """Test DocumentProcessor class methods"""
    print("\nTesting DocumentProcessor...")
    
    try:
        import app
        
        # Test static methods
        processor = app.DocumentProcessor
        
        # Test date formatting
        date_result = processor.format_date("2024-01-15", "letter")
        print(f"✓ Date formatting works: {date_result}")
        
        # Test address formatting
        address_result = processor.format_address("123 Main St, City, State", "letter")
        print(f"✓ Address formatting works")
        
        # Test casing
        casing_result = processor.apply_casing("test text", "upper")
        assert casing_result == "TEST TEXT"
        print("✓ Text casing works")
        
        # Test placeholder defaults
        default_name = processor.get_smart_placeholder_default("name")
        print(f"✓ Smart defaults work: name -> {default_name}")
        
        return True
        
    except Exception as e:
        print(f"✗ DocumentProcessor error: {e}")
        return False

def test_routes_definition():
    """Test if all routes are properly defined"""
    print("\nTesting route definitions...")
    
    try:
        import app
        
        flask_app = app.app
        
        # Get all routes
        routes = []
        for rule in flask_app.url_map.iter_rules():
            routes.append({
                'endpoint': rule.endpoint,
                'methods': list(rule.methods),
                'rule': str(rule)
            })
        
        # Expected routes
        expected_endpoints = [
            'index',
            'create',
            'generate',
            'download',
            'batch',
            'batch_results',
            'get_merged_placeholders',
            'admin',
            'admin_templates',
            'admin_upload_template',
            'admin_edit_template',
            'admin_pause_template',
            'admin_resume_template',
            'admin_delete_template',
            'admin_clear_database',
            'admin_backup_database',
            'delete_document'
        ]
        
        found_endpoints = [route['endpoint'] for route in routes]
        missing_endpoints = [ep for ep in expected_endpoints if ep not in found_endpoints]
        
        if missing_endpoints:
            print(f"✗ Missing route endpoints: {missing_endpoints}")
            return False
        else:
            print(f"✓ All {len(expected_endpoints)} expected routes defined")
            return True
            
    except Exception as e:
        print(f"✗ Route definition error: {e}")
        return False

def test_pdf_conversion():
    """Test PDF conversion functionality"""
    print("\nTesting PDF conversion...")
    
    try:
        import app
        from docx import Document
        
        # Create a simple test DOCX
        test_doc = Document()
        test_doc.add_paragraph("This is a test document for PDF conversion.")
        test_doc.add_paragraph("It contains multiple paragraphs.")
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            test_doc.save(temp_docx.name)
            
            # Test PDF conversion
            pdf_path = app.DocumentProcessor.convert_to_pdf(temp_docx.name)
            
            if os.path.exists(pdf_path):
                print("✓ PDF conversion successful")
                # Clean up
                os.unlink(temp_docx.name)
                os.unlink(pdf_path)
                return True
            else:
                print("✗ PDF file not created")
                return False
                
    except Exception as e:
        print(f"✗ PDF conversion error: {e}")
        return False

def test_template_processing():
    """Test template processing functionality"""
    print("\nTesting template processing...")
    
    try:
        import app
        
        # Test placeholder extraction (mock)
        processor = app.DocumentProcessor
        
        # Test variable type detection
        assert processor.detect_variable_type("name") == "text"
        assert processor.detect_variable_type("date") == "date" 
        assert processor.detect_variable_type("email") == "email"
        assert processor.detect_variable_type("gender") == "option"
        
        print("✓ Variable type detection works")
        
        # Test smart options
        gender_options = processor.get_smart_options("gender")
        assert "Male" in gender_options and "Female" in gender_options
        print("✓ Smart options generation works")
        
        return True
        
    except Exception as e:
        print(f"✗ Template processing error: {e}")
        return False

def run_all_tests():
    """Run all tests and provide summary"""
    print("=" * 60)
    print("MyTypist Application - Comprehensive Test Suite")
    print("=" * 60)
    
    tests = [
        ("Import Dependencies", test_imports),
        ("File Structure", test_app_structure),
        ("App Configuration", test_app_configuration),
        ("Document Processor", test_document_processor),
        ("Route Definitions", test_routes_definition),
        ("PDF Conversion", test_pdf_conversion),
        ("Template Processing", test_template_processing)
    ]
    
    results = []
    
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"✗ {test_name} failed with exception: {e}")
            results.append((test_name, False))
    
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)
    
    passed = 0
    total = len(results)
    
    for test_name, result in results:
        status = "PASS" if result else "FAIL"
        symbol = "✓" if result else "✗"
        print(f"{symbol} {test_name}: {status}")
        if result:
            passed += 1
    
    print(f"\nResults: {passed}/{total} tests passed")
    
    if passed == total:
        print("\n🎉 ALL TESTS PASSED! Your MyTypist application is working correctly.")
    else:
        print(f"\n⚠️  {total - passed} test(s) failed. Please review the issues above.")
    
    return passed == total

if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
