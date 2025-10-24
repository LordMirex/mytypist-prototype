"""
Test script to verify PDF conversion using Aspose.Words
"""
import os
import logging
from docxtpl import DocxTemplate
import aspose.words as aw

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_test_docx():
    """Create a test DOCX file with various formatting"""
    logger.info("Creating test DOCX file...")
    
    # Create a new DOCX file from scratch
    from docx import Document
    doc = Document()
    
    # Add some test content with various formatting
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a normal paragraph.')
    
    # Add styled paragraph
    p = doc.add_paragraph()
    p.add_run('This is bold text.').bold = True
    p.add_run(' And this is normal text.')
    p.add_run(' This is italic.').italic = True
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.text = 'Table cell'
            
    output_docx = "test_output.docx"
    doc.save(output_docx)
    
    if os.path.exists(output_docx):
        logger.info(f"DOCX generated successfully: {output_docx}")
        return output_docx
    else:
        raise RuntimeError("Failed to create test DOCX")
    context = {
        'title': 'Test Document',
        'content': 'This is a test document with some content.',
        'footer': 'Page 1'
    }
    
    output_docx = "test_output.docx"
    doc.render(context)
    doc.save(output_docx)
    
    if os.path.exists(output_docx):
        logger.info(f"DOCX generated successfully: {output_docx}")
        return output_docx
    else:
        raise RuntimeError("Failed to create test DOCX")

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF using Aspose.Words"""
    logger.info(f"PDF export started via Aspose.Words: {docx_path}")
    
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        # Set globalization mode to invariant
        os.environ["System.Globalization.Invariant"] = "true"
        
        # Load DOCX and save as PDF
        doc = aw.Document(docx_path)
        doc.save(pdf_path)
        
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            logger.info(f"PDF export complete → {pdf_path}")
            return True, pdf_path
        else:
            logger.error("PDF conversion failed - output file was not created or is empty")
            return False, None
    except Exception as e:
        logger.error(f"PDF conversion failed: {str(e)}")
        return False, None

def cleanup(files):
    """Clean up test files"""
    for f in files:
        if os.path.exists(f):
            os.remove(f)
            logger.info(f"Cleaned up: {f}")

def main():
    """Run the conversion test"""
    test_files = []
    try:
        # Create test DOCX
        docx_path = create_test_docx()
        test_files.append(docx_path)
        
        # Convert to PDF
        success, pdf_path = convert_to_pdf(docx_path)
        if pdf_path:
            test_files.append(pdf_path)
        
        if success:
            logger.info("✅ Conversion test passed successfully!")
            logger.info(f"PDF output: {pdf_path}")
        else:
            logger.error("❌ Conversion test failed!")
    finally:
        cleanup(test_files)

if __name__ == '__main__':
    main()